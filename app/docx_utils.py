"""
Utilities to extract text from docx and to write highlighted review + review table.
"""
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

def extract_text(path: str) -> str:
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paras.append(text)
    # include tables text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([c.text.strip() for c in row.cells])
            if row_text.strip():
                paras.append(row_text)
    return "\n".join(paras)

def highlight_runs_for_phrase(doc: Document, phrase: str):
    phrase_lower = phrase.lower()
    for p in doc.paragraphs:
        for run in p.runs:
            if phrase_lower in run.text.lower():
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def append_review_table(doc: Document, issues: list):
    if not issues:
        doc.add_paragraph("Review: No issues found.")
        return
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Document"
    hdr[1].text = "Section"
    hdr[2].text = "Issue"
    hdr[3].text = "Suggestion"
    for it in issues:
        r = table.add_row().cells
        r[0].text = it.get("document", "-")
        r[1].text = it.get("section", "-")
        r[2].text = it.get("issue", "-")
        r[3].text = it.get("suggestion", "-")

def save_reviewed_docx(original_path: str, out_path: str, issues: list, highlights: list = None):
    doc = Document(original_path)
    # apply simple highlights for each phrase
    if highlights:
        for phrase in highlights:
            highlight_runs_for_phrase(doc, phrase)
    # append review table
    append_review_table(doc, issues)
    doc.save(out_path)

if __name__ == "__main__":
    p = Path(__file__).parent / ".." / "sample_before.docx"
    if p.exists():
        txt = extract_text(str(p))
        print(txt[:400])